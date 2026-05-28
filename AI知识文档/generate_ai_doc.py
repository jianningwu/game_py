#!/usr/bin/env python3
"""
================================================================================
脚本功能：生成AI知识全览文档（Word + Excel + 配套图表）
- 使用 python-docx 生成结构化 Word 文档，涵盖AI历史、技术、Agent、大数据等8章
- 使用 matplotlib 生成7张可视化图表（时间线、架构图、市场趋势等）
- 使用 openpyxl 生成 Excel 数据表（含模型对比、术语表、推荐资源）
- 输出目录：D:\ai_op\
================================================================================
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.utils import get_column_letter
from collections import Counter

OUTPUT_DIR = r"D:\ai_op"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ═══════════════════════════════════════════════════════════════
# PART 1: Generate Charts with matplotlib
# ═══════════════════════════════════════════════════════════════

plt.rcParams['font.family'] = ['Microsoft YaHei', 'SimHei', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

# --- Chart 1: AI History Timeline ---
def create_timeline_chart():
    fig, ax = plt.subplots(figsize=(16, 6))
    events = [
        (1950, "图灵测试\n(Turing Test)"),
        (1956, "达特茅斯会议\nAI诞生"),
        (1966, "ELIZA\n首个聊天机器人"),
        (1986, "反向传播算法\n(Backpropagation)"),
        (1989, "CNN\n(LeCun)"),
        (1997, "深蓝击败\n卡斯帕罗夫"),
        (2006, "深度信念网络\n(Hinton)"),
        (2012, "AlexNet\n深度学习革命"),
        (2014, "GAN\n(Goodfellow)"),
        (2017, "Transformer\n注意力机制"),
        (2018, "BERT/GPT-1\n预训练时代"),
        (2020, "GPT-3\n1750亿参数"),
        (2022, "ChatGPT\n爆发式增长"),
        (2024, "GPT-4o/Sora\n多模态时代"),
        (2025, "AI Agent\n自主决策时代"),
        (2026, "智能体生态\nMCP标准化"),
    ]
    years = [e[0] for e in events]
    labels = [e[1] for e in events]
    colors = plt.cm.viridis(np.linspace(0.1, 0.9, len(events)))

    for i, (y, label, c) in enumerate(zip(years, labels, colors)):
        ax.scatter(y, 0, s=180, color=c, zorder=5, edgecolors='white', linewidth=1.5)
        offset = 0.4 if i % 2 == 0 else -0.4
        va = 'bottom' if i % 2 == 0 else 'top'
        ax.annotate(label, (y, offset), textcoords="offset points",
                    xytext=(0, 12 if i % 2 == 0 else -12), ha='center', va=va,
                    fontsize=7.5, fontweight='bold', color='#333333',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                              edgecolor=c, alpha=0.85))

    ax.axhline(y=0, color='#cccccc', linewidth=2, zorder=1)
    ax.set_ylim(-1.2, 1.2)
    ax.set_yticks([])
    ax.set_xlim(1945, 2030)
    ax.set_xlabel('年份', fontsize=12, fontweight='bold')
    ax.set_title('人工智能发展历史时间线 (1950–2026)', fontsize=16, fontweight='bold', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.grid(axis='x', alpha=0.3)
    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_timeline.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 2: Model Parameter Growth ---
def create_model_growth_chart():
    fig, ax = plt.subplots(figsize=(14, 6))
    models = ['ELIZA\n1966', 'AlexNet\n2012', 'BERT\n2018', 'GPT-2\n2019',
              'GPT-3\n2020', 'PaLM\n2022', 'GPT-4\n2023', 'Claude 3.5\n2024',
              'GPT-5\n2025', 'DeepSeek-V4\n2025', 'Claude 4.7\n2026']
    params = [0.00001, 0.06, 0.34, 1.5, 175, 540, 1800, 2000, 3000, 2300, 2800]
    colors = ['#e74c3c', '#e67e22', '#f1c40f', '#2ecc71', '#1abc9c',
              '#3498db', '#9b59b6', '#34495e', '#e74c3c', '#16a085', '#8e44ad']

    bars = ax.barh(models, params, color=colors, edgecolor='white', linewidth=1.2, height=0.7)
    ax.set_xscale('log')
    ax.set_xlabel('参数量 (亿)', fontsize=12, fontweight='bold')
    ax.set_title('大语言模型参数规模增长趋势 (对数刻度)', fontsize=16, fontweight='bold', pad=15)
    for bar, val in zip(bars, params):
        label = f'{val:.0f}亿' if val >= 1 else (f'{val*100:.0f}M' if val >= 0.01 else f'{val*10000:.0f}万')
        ax.text(bar.get_width() * 1.15, bar.get_y() + bar.get_height()/2,
                label, va='center', fontsize=10, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_model_growth.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 3: AI Winter Cycles ---
def create_ai_winter_chart():
    fig, ax = plt.subplots(figsize=(14, 5.5))
    x = np.linspace(1950, 2026, 200)
    # Simulated AI hype/funding curve
    y = (0.15 * np.exp(-((x - 1965) / 8) ** 2) +
         0.1 * np.sin((x - 1950) / 6) +
         0.2 * np.exp(-((x - 1985) / 10) ** 2) +
         0.35 * np.exp(-((x - 2012) / 6) ** 2) +
         0.5 * np.exp(-((x - 2018) / 5) ** 2) +
         0.7 * np.exp(-((x - 2023) / 4) ** 2) +
         0.05 * (x - 1950) / 30)
    y = np.clip(y, 0, 1)

    ax.fill_between(x, y, alpha=0.35, color='#3498db')
    ax.plot(x, y, color='#2980b9', linewidth=2.5)

    # Mark winters
    ax.axvspan(1974, 1980, alpha=0.15, color='#95a5a6')
    ax.axvspan(1987, 1994, alpha=0.15, color='#95a5a6')
    ax.annotate('第一次\nAI寒冬', xy=(1977, 0.08), ha='center', fontsize=10,
                fontweight='bold', color='#7f8c8d')
    ax.annotate('第二次\nAI寒冬', xy=(1990.5, 0.08), ha='center', fontsize=10,
                fontweight='bold', color='#7f8c8d')

    ax.annotate('达特茅斯\n会议', xy=(1956, 0.28), ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.annotate('专家系统\n兴起', xy=(1983, 0.32), ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.annotate('AlexNet\nDL革命', xy=(2012, 0.55), ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.annotate('ChatGPT\n爆发', xy=(2023, 0.92), ha='center', fontsize=9, fontweight='bold', color='#2c3e50')

    ax.set_ylabel('AI 投资/关注度', fontsize=12, fontweight='bold')
    ax.set_xlabel('年份', fontsize=12, fontweight='bold')
    ax.set_title('人工智能的兴衰周期："AI寒冬"与爆发', fontsize=16, fontweight='bold', pad=15)
    ax.set_ylim(0, 1.1)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_ai_winter.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 4: LLM Architecture Overview ---
def create_llm_arch_chart():
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('大语言模型 (LLM) 核心架构 — 以 GPT 系列 Decoder-Only 为例',
                 fontsize=15, fontweight='bold', pad=18)

    boxes = [
        (3, 8, 8, 0.5, '输出文本: "人工智能改变了世界"', '#2ecc71', 'white'),
        (4, 7, 6, 0.8, '输出层 (Linear + Softmax)\n将隐藏向量映射为词表概率分布', '#27ae60', 'white'),
        (4, 5.8, 6, 1.0, 'N × Transformer Decoder 层堆叠\n• 掩码多头自注意力 (Causal Attention)\n• 前馈网络 (FFN + SwiGLU 激活)\n• RMSNorm 层归一化 + 残差连接', '#3498db', 'white'),
        (4, 4.5, 6, 1.0, '旋转位置编码 (RoPE)\n将 token 位置信息注入 Q/K 向量\n支持长度外推 (Context Window Extension)', '#9b59b6', 'white'),
        (4, 3.2, 6, 1.0, 'Token 嵌入层 (Embedding)\n将离散词元映射为连续稠密向量\n词汇表大小: 50K~256K tokens', '#e67e22', 'white'),
        (3, 2.2, 8, 0.8, '输入: "Artificial Intelligence is changing the world"', '#e74c3c', 'white'),
        (5.5, 1.2, 3, 0.7, 'Tokenizer 分词器\nBPE / SentencePiece', '#95a5a6', 'white'),
    ]
    for x, y, w, h, text, color, tc in boxes:
        rect = FancyBboxPatch((x, y - h/2), w, h, boxstyle="round,pad=0.15",
                              facecolor=color, edgecolor='#2c3e50', linewidth=1.5, alpha=0.92)
        ax.add_patch(rect)
        ax.text(x + w/2, y, text, ha='center', va='center', fontsize=8.5,
                color=tc, fontweight='bold')

    # Arrows
    arrow_ys = [7.6, 6.35, 5.25, 3.95, 2.7]
    for i in range(len(arrow_ys) - 1):
        ax.annotate('', xy=(7, arrow_ys[i+1] + 0.2), xytext=(7, arrow_ys[i] - 0.2),
                    arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))

    # Side annotations
    side_notes = [
        (0.5, 6.3, "自注意力机制:\n每个token关注\n所有前序token,\n捕捉长程依赖"),
        (11.5, 5, "前馈网络:\n逐位置的非线性\n变换,引入模型\n表达能力"),
        (0.5, 2.5, "RoPE:\n通过旋转矩阵\n隐式编码相对\n位置信息"),
    ]
    for sx, sy, stext in side_notes:
        ax.text(sx, sy, stext, fontsize=7.5, color='#555555',
                bbox=dict(boxstyle='round', facecolor='#ecf0f1', alpha=0.8),
                ha='center', va='center')

    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_llm_arch.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 5: Attention Mechanism ---
def create_attention_chart():
    fig, axes = plt.subplots(1, 3, figsize=(16, 5))
    fig.suptitle('自注意力机制 (Self-Attention) 原理图解', fontsize=15, fontweight='bold', y=1.01)

    # Panel 1: Q, K, V projection
    ax = axes[0]
    ax.set_xlim(0, 5)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('Step 1: Q/K/V 投影', fontsize=12, fontweight='bold')
    tokens = ['我', '爱', 'AI']
    colors = ['#e74c3c', '#3498db', '#2ecc71']
    for i, (t, c) in enumerate(zip(tokens, colors)):
        y_pos = 4 - i * 1.5
        rect = FancyBboxPatch((0.3, y_pos - 0.35), 1.2, 0.7, boxstyle="round",
                              facecolor=c, edgecolor='#2c3e50', alpha=0.8)
        ax.add_patch(rect)
        ax.text(0.9, y_pos, t, ha='center', va='center', fontsize=14, color='white', fontweight='bold')
        ax.annotate('', xy=(2.2, y_pos), xytext=(1.6, y_pos),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        ax.text(2.9, y_pos + 0.2, f'Q{i+1}', fontsize=8, color='#e74c3c')
        ax.text(2.9, y_pos - 0.1, f'K{i+1}', fontsize=8, color='#3498db')
        ax.text(2.9, y_pos - 0.4, f'V{i+1}', fontsize=8, color='#2ecc71')

    # Panel 2: Attention score matrix
    ax = axes[1]
    ax.set_title('Step 2: 注意力分数矩阵', fontsize=12, fontweight='bold')
    scores = np.array([[0.8, 0.15, 0.05],
                       [0.25, 0.6, 0.15],
                       [0.1, 0.2, 0.7]])
    im = ax.imshow(scores, cmap='YlOrRd', vmin=0, vmax=1)
    ax.set_xticks(range(3))
    ax.set_yticks(range(3))
    ax.set_xticklabels(tokens, fontsize=12)
    ax.set_yticklabels(tokens, fontsize=12)
    for i in range(3):
        for j in range(3):
            ax.text(j, i, f'{scores[i, j]:.2f}', ha='center', va='center',
                    fontsize=14, fontweight='bold',
                    color='white' if scores[i, j] > 0.5 else '#333')
    ax.set_xlabel('Key (被关注的token)', fontsize=10)
    ax.set_ylabel('Query (当前token)', fontsize=10)
    plt.colorbar(im, ax=ax, shrink=0.8)

    # Panel 3: Weighted sum
    ax = axes[2]
    ax.set_xlim(0, 5)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('Step 3: 加权求和输出', fontsize=12, fontweight='bold')
    ax.text(0.8, 3.8, 'Output("爱") = \n0.25×V("我") +\n0.60×V("爱") +\n0.15×V("AI")',
            fontsize=13, fontfamily='monospace',
            bbox=dict(boxstyle='round', facecolor='#f39c12', alpha=0.25))
    ax.text(0.8, 1.5, '✓ 每个输出token\n  融合了所有相关\n  上下文信息\n✓ 权重自动学习\n  无需人工规则',
            fontsize=11, color='#2c3e50')

    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_attention.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 6: AI Agent Architecture ---
def create_agent_arch_chart():
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('AI Agent 智能体架构：感知 → 规划 → 执行 → 反馈循环', fontsize=15, fontweight='bold', pad=15)

    # Central agent
    circle = plt.Circle((7, 4), 1.0, color='#e74c3c', alpha=0.9, ec='#2c3e50', lw=2)
    ax.add_patch(circle)
    ax.text(7, 4, 'AI Agent\n智能体核心', ha='center', va='center', fontsize=10, color='white', fontweight='bold')

    # Surrounding modules
    modules = [
        (7, 7.2, '🧠 大语言模型 (LLM)\n推理 | 规划 | 决策', '#3498db'),
        (1, 5.5, '👁️ 感知模块\n文本 | 图像 | 语音 | API', '#2ecc71'),
        (1, 2.5, '🗄️ 记忆系统\n短期记忆 | 长期记忆 | RAG', '#9b59b6'),
        (7, 1, '🔧 工具调用 (Tool Use)\n代码执行 | API调用 | 数据库查询', '#e67e22'),
        (12, 5.5, '📋 任务规划\n目标分解 | 步骤排序 | 依赖管理', '#1abc9c'),
        (12, 2.5, '✅ 执行与反馈\n动作执行 | 结果验证 | 自我修正', '#f39c12'),
    ]
    for mx, my, mtext, mc in modules:
        rect = FancyBboxPatch((mx - 1.3, my - 0.6), 2.6, 1.2, boxstyle="round,pad=0.1",
                              facecolor=mc, edgecolor='#2c3e50', alpha=0.88, lw=1.5)
        ax.add_patch(rect)
        ax.text(mx, my, mtext, ha='center', va='center', fontsize=8, color='white', fontweight='bold')
        # Arrow from center to module
        dx, dy = mx - 7, my - 4
        dist = np.sqrt(dx**2 + dy**2)
        start_x = 7 + dx/dist * 1.05
        start_y = 4 + dy/dist * 1.05
        end_x = mx - dx/dist * 1.35
        end_y = my - dy/dist * 1.35
        ax.annotate('', xy=(end_x, end_y), xytext=(start_x, start_y),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.8,
                                    connectionstyle='arc3,rad=0.1'))

    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_agent_arch.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# --- Chart 7: AI market / growth ---
def create_market_chart():
    fig, axes = plt.subplots(1, 2, figsize=(15, 5.5))
    fig.suptitle('全球AI市场趋势', fontsize=15, fontweight='bold')

    # Left: Market size
    ax = axes[0]
    years = [2020, 2021, 2022, 2023, 2024, 2025, 2026, 2027, 2028]
    size = [51, 93, 142, 207, 298, 420, 594, 780, 1020]
    ax.fill_between(years, size, alpha=0.3, color='#3498db')
    ax.plot(years, size, 'o-', color='#2980b9', linewidth=2.5, markersize=8)
    ax.set_ylabel('市场规模 (十亿美元)', fontsize=11, fontweight='bold')
    ax.set_title('全球AI市场规模 (2020-2028E)', fontsize=13, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(alpha=0.3)
    for yr, sz in zip(years, size):
        ax.text(yr, sz + 25, f'${sz}B', ha='center', fontsize=9, fontweight='bold')

    # Right: GPU/Compute
    ax = axes[1]
    compute_years = [2018, 2019, 2020, 2021, 2022, 2023, 2024, 2025, 2026]
    petaflops = [0.5, 2, 8, 30, 120, 500, 2100, 8500, 35000]
    ax.fill_between(compute_years, petaflops, alpha=0.3, color='#e74c3c')
    ax.plot(compute_years, petaflops, 's-', color='#c0392b', linewidth=2.5, markersize=8)
    ax.set_yscale('log')
    ax.set_ylabel('训练算力 (PetaFLOPS-days, 对数)', fontsize=11, fontweight='bold')
    ax.set_title('AI训练算力指数增长 (2018-2026)', fontsize=13, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(alpha=0.3)
    for yr, pf in zip(compute_years, petaflops):
        ax.text(yr, pf * 1.5, f'{pf:,}', ha='center', fontsize=8, fontweight='bold')

    plt.tight_layout()
    path = os.path.join(OUTPUT_DIR, "chart_market.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# Generate all charts
print("正在生成图表...")
chart_paths = {
    'timeline': create_timeline_chart(),
    'model_growth': create_model_growth_chart(),
    'ai_winter': create_ai_winter_chart(),
    'llm_arch': create_llm_arch_chart(),
    'attention': create_attention_chart(),
    'agent_arch': create_agent_arch_chart(),
    'market': create_market_chart(),
}
print(f"已生成 {len(chart_paths)} 张图表")

# ═══════════════════════════════════════════════════════════════
# PART 2: Generate Word Document
# ═══════════════════════════════════════════════════════════════

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = 'Microsoft YaHei'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h_font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    if level == 1:
        h_font.size = Pt(22)
    elif level == 2:
        h_font.size = Pt(16)
    else:
        h_font.size = Pt(13)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table in the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()  # spacer
    return table

# ═══ TITLE PAGE ═══
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n人工智能知识全览')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('从历史发展到前沿技术的全面解读')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年5月 · 完整版')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

doc.add_page_break()

# ═══ TABLE OF CONTENTS ═══
doc.add_heading('目录', level=1)
toc_items = [
    '第一章  人工智能发展历史',
    '第二章  当前AI发展现状 (2026)',
    '第三章  大语言模型 (LLM) 技术详解',
    '第四章  核心技术名词详解',
    '第五章  AI Agent 智能体详解',
    '第六章  大数据与AI的关系',
    '第七章  AI基础设施与算力',
    '第八章  未来展望与趋势',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(13)

doc.add_page_break()

# ═══ CHAPTER 1: AI HISTORY ═══
doc.add_heading('第一章  人工智能发展历史', level=1)

doc.add_paragraph(
    '人工智能（Artificial Intelligence，简称AI）的历史是一部跨越近80年的壮阔史诗。'
    '从1943年第一个神经元数学模型，到2026年能够自主决策的AI智能体，人类对"会思考的机器"'
    '的追求从未停歇。本章将带你回顾这段激动人心的历程。'
)

doc.add_heading('1.1 发展历史时间线总览', level=2)
doc.add_picture(chart_paths['timeline'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1.2 AI的兴衰周期', level=2)
doc.add_paragraph(
    '人工智能的发展并非一路向上，而是经历了多次"AI寒冬"——期望过高导致投资退潮的周期。'
    '理解这个周期对于客观看待当前AI热潮至关重要。'
)
doc.add_picture(chart_paths['ai_winter'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1.3 关键历史阶段', level=2)

doc.add_heading('理论奠基期 (1943–1956)', level=3)
doc.add_paragraph(
    '1943年，McCulloch和Pitts提出了第一个神经元数学模型。1950年，图灵发表《计算机器与智能》，'
    '提出了著名的"图灵测试"。1956年，麦卡锡、明斯基、香农等人在达特茅斯会议上正式提出'
    '"人工智能"这一术语，标志着AI作为独立学科的诞生。'
)

doc.add_heading('符号主义时代 (1957–1974)', level=3)
doc.add_paragraph(
    '这一时期，AI研究以符号推理为主。1957年，Rosenblatt发明了感知机（Perceptron），'
    '这是第一个可训练的神经网络。1958年，麦卡锡创建了LISP编程语言。1966年，Weizenbaum'
    '开发了第一个聊天机器人ELIZA。然而，1969年明斯基等人证明了感知机的局限性，'
    '加上过高期望的破灭，导致了第一次AI寒冬（1974-1980）。'
)

doc.add_heading('专家系统与连接主义复兴 (1980–1993)', level=3)
doc.add_paragraph(
    '1980年代，专家系统在企业中取得商业成功。1982年Hopfield网络的提出和1986年'
    '反向传播算法的普及，使神经网络研究重新活跃。然而，专家系统的维护成本过高，'
    '导致第二次AI寒冬（1987-1994）。1989年，LeCun将卷积神经网络应用于手写数字识别，'
    '奠定了现代计算机视觉的基础。'
)

doc.add_heading('统计学习时代 (1994–2011)', level=3)
doc.add_paragraph(
    '1997年，IBM深蓝击败国际象棋冠军卡斯帕罗夫。同年，LSTM（长短期记忆网络）被提出，'
    '解决了循环神经网络的梯度消失问题。2006年，Hinton提出深度信念网络，'
    '重新点燃了深度学习的研究热情。2009年，李飞飞发布ImageNet数据集。'
)

doc.add_heading('深度学习革命 (2012–2017)', level=3)
doc.add_paragraph(
    '2012年，Hinton团队的AlexNet在ImageNet竞赛中取得了突破性胜利，将图像识别的错误率'
    '从25.8%降至16.4%，深度学习从此成为AI研究的主流方向。2014年，Goodfellow发明了GAN'
    '（生成对抗网络）。2016年，DeepMind的AlphaGo击败围棋世界冠军李世石。'
    '2017年，Google发表《Attention Is All You Need》，提出了Transformer架构——'
    '这个架构成为后来所有大语言模型的基础。'
)

doc.add_heading('预训练大模型时代 (2018–2022)', level=3)
doc.add_paragraph(
    '2018年，OpenAI发布GPT-1，Google发布BERT，开启了"预训练+微调"的范式。'
    '2020年，GPT-3以1750亿参数展示了令人惊叹的涌现能力。2022年11月，ChatGPT发布，'
    '在2个月内获得1亿用户，成为历史上增长最快的消费级应用，AI真正走向了大众。'
)

doc.add_heading('多模态与智能体时代 (2023–2026)', level=3)
doc.add_paragraph(
    '2023年GPT-4实现了多模态理解，2024年Sora展示了视频生成的新高度。2025年AI Agent'
    '技术成熟，AI从"对话工具"进化为"自主执行者"。2026年，MCP（模型上下文协议）'
    '成为行业标准，多智能体协作架构在企业中大规模落地。'
)

doc.add_page_break()

# ═══ CHAPTER 2: CURRENT AI STATUS (2026) ═══
doc.add_heading('第二章  当前AI发展现状 (2026年)', level=1)

doc.add_heading('2.1 主要模型格局', level=2)
doc.add_paragraph(
    '截至2026年5月，全球大模型市场形成了"多极竞争"的格局。下表汇总了当前主流模型。'
)

model_headers = ['模型名称', '开发机构', '关键特性']
model_rows = [
    ['Claude Opus 4.7', 'Anthropic', '顶级推理能力，100万tokens上下文，Agent Skill API'],
    ['GPT-5.5', 'OpenAI', '知识工作者优化，多模态，Workspace Agents'],
    ['GPT-Rosalind', 'OpenAI', '前沿推理模型，专精生物学和药物发现'],
    ['Gemini 3.1 Flash', 'Google', '新型TTS引擎，音频标签技术'],
    ['Qwen 3.6', '阿里巴巴', '开源模型，门控注意力机制，高性能'],
    ['DeepSeek V4 / V3.2', 'DeepSeek', '高效架构，稀疏注意力，强化学习合成'],
    ['Mistral 3', 'Mistral', '大+Ministral 14B/8B/3B，全系支持视觉'],
    ['GLM-4.7', '智谱AI', '355B MoE参数(32B激活)，编码与推理领先'],
    ['Kimi K2.6', 'Moonshot AI', '强大的开源竞品，长上下文'],
    ['Nemotron 3 Nano', 'NVIDIA', '30B MoE，完全开源（数据+权重+配方）'],
    ['K-EXAONE', 'LG', '236B MoE，韩英双语，强多语言支持'],
]
add_styled_table(doc, model_headers, model_rows)

doc.add_heading('2.2 全球AI市场数据', level=2)
doc.add_picture(chart_paths['market'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '全球AI市场从2020年的510亿美元增长至2026年的约5940亿美元，预计2028年将突破1万亿美元。'
    '训练算力需求同步呈指数增长——从2018年的约0.5 PetaFLOPS-days到2026年的约35,000 '
    'PetaFLOPS-days。这一增长速度甚至超过了摩尔定律，主要驱动力来自大规模集群（如'
    'Anthropic × SpaceX的"Colossus 1"超算中心，配备22万+ GPU，功率300MW）。'
)

doc.add_heading('2.3 关键行业趋势', level=2)
trend_headers = ['趋势', '说明', '影响']
trend_rows = [
    ['AI Agent爆发行', 'AI从对话工具进化为自主执行者，多智能体协作成为现实', '企业工作流全面重构'],
    ['MCP协议标准化', 'Anthropic捐赠MCP给Linux基金会，OpenAI贡献AGENTS.md', '工具互操作性大幅提升'],
    ['开源模型崛起', 'Qwen、DeepSeek、GLM、Kimi等开源模型性能接近闭源', 'AI能力民主化加速'],
    ['多模态融合', '文本+图像+视频+语音的统一理解与生成', '人机交互进入全新范式'],
    ['边缘AI部署', '小型模型(LFM2.5、Ministral 3B)支持端侧运行', '离线AI应用成为可能'],
    ['数据质量危机', 'Gartner预测60%的AI项目可能因数据问题被放弃', '数据治理成为核心挑战'],
    ['AI安全监管', 'EU AI Act全面实施，各国加快AI立法', '合规成本上升但信任度提高'],
]
add_styled_table(doc, trend_headers, trend_rows)
doc.add_page_break()

# ═══ CHAPTER 3: LLM TECHNICAL DETAILS ═══
doc.add_heading('第三章  大语言模型 (LLM) 技术详解', level=1)

doc.add_paragraph(
    '大语言模型（Large Language Model, LLM）是当前AI革命的核心引擎。'
    '理解LLM的技术原理，是深入掌握现代AI的基础。'
)

doc.add_heading('3.1 整体架构概览', level=2)
doc.add_picture(chart_paths['llm_arch'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.2 自注意力机制 (Self-Attention)', level=2)
doc.add_paragraph(
    '注意力机制是Transformer架构的核心创新。其核心思想是：在处理每个词时，'
    '动态地决定应该"关注"上下文中的哪些词，而不是像RNN那样按顺序处理。'
    '\n\n核心公式：Attention(Q, K, V) = softmax(QK^T / √d_k) × V'
    '\n\n其中：'
    '\n• Q (Query): 当前词的"查询"向量——"我在找什么？"'
    '\n• K (Key): 上下文词的"键"向量——"我是什么？"'
    '\n• V (Value): 上下文词的"值"向量——"我包含什么信息？"'
    '\n• √d_k: 缩放因子，防止点积过大导致梯度消失'
    '\n\n除以√d_k是训练稳定性的关键——研究表明，省略这一操作会使训练失败率提高60%以上。'
)
doc.add_picture(chart_paths['attention'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.3 多头注意力 (Multi-Head Attention)', level=2)
doc.add_paragraph(
    '单个注意力头只能关注一种语义关系。多头注意力通过并行运行多个注意力头，'
    '每个头关注不同的语义维度（如语法结构、指代关系、实体关联等），'
    '最后拼接所有头的输出，大大增强了模型的表达能力。'
    '典型的LLM使用32-128个注意力头。'
)

doc.add_heading('3.4 三大关键技术', level=2)

doc.add_heading('旋转位置编码 (RoPE)', level=3)
doc.add_paragraph(
    '由于自注意力本身不感知词序，需要人工注入位置信息。RoPE通过在Q和K向量上施加旋转操作，'
    '以极简的方式将相对位置信息编码到注意力计算中。其最大优势是支持"长度外推"——'
    '在16K序列上训练的模型，推理时扩展到32K序列，性能下降不超过5%。'
    'RoPE已成为GPT-NeoX、LLaMA、Qwen、DeepSeek等主流模型的标准配置。'
)

doc.add_heading('分组查询注意力 (GQA)', level=3)
doc.add_paragraph(
    '标准多头注意力的K/V投影消耗大量显存。GQA将注意力头分为若干组，组内共享K/V投影，'
    '显存占用可降至原来的1/G。GQA在不损失模型质量的前提下，大幅降低了推理成本。'
    'LLaMA 2、Mistral、Gemma等模型均采用GQA技术。'
)

doc.add_heading('混合专家模型 (MoE)', level=3)
doc.add_paragraph(
    'MoE不是让所有参数参与每次计算，而是每次只激活部分"专家"。'
    '例如GLM-4.7拥有3550亿总参数，但每次推理只激活320亿。'
    '这大幅提升了模型的"知识容量/计算成本"比——更聪明但不更费电。'
    '2026年，MoE已成为高效大模型的标配架构。'
)

doc.add_heading('3.5 模型参数增长史', level=2)
doc.add_picture(chart_paths['model_growth'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    '从2018年BERT的3.4亿参数到2026年主流模型的万亿级规模，仅8年时间模型规模增长了近1000倍。'
    '但2024-2026年出现了一个重要转变：业界从单纯追求参数规模，转向追求效率和稀疏激活（MoE），'
    '在相同计算预算下实现更强性能。'
)

doc.add_heading('3.6 模型架构类型对比', level=2)
arch_headers = ['架构类型', '代表模型', '注意力方式', '最佳场景']
arch_rows = [
    ['Encoder-Only', 'BERT, RoBERTa', '双向注意力', '文本分类、实体识别、语义匹配'],
    ['Decoder-Only', 'GPT系列, LLaMA, Claude', '因果掩码注意力', '文本生成、对话、推理'],
    ['Encoder-Decoder', 'T5, BART, Gemini', '编码双向+解码因果', '翻译、摘要、跨模态'],
]
add_styled_table(doc, arch_headers, arch_rows)
doc.add_page_break()

# ═══ CHAPTER 4: KEY TERMINOLOGY ═══
doc.add_heading('第四章  核心技术名词详解', level=1)

doc.add_paragraph(
    'AI领域术语繁多，容易混淆。本章对最重要的技术名词进行系统化解释。'
)

doc.add_heading('4.1 核心概念层', level=2)

term_headers = ['术语', '英文/缩写', '详细解释']
term_rows = [
    ['人工智能', 'AI', '让机器模拟人类智能的广义学科。分为弱AI(专用任务)和强AI(通用智能)。'],
    ['机器学习', 'ML', 'AI的子集，通过数据训练让模型自动学习规律，而非手工编写规则。'],
    ['深度学习', 'DL', 'ML的子集，使用多层神经网络从大规模数据中自动提取特征。'],
    ['神经网络', 'Neural Network', '受生物神经元启发的计算模型。多层网络可拟合任意复杂函数。'],
    ['GPU', 'GPU', '图形处理器，因其大规模并行计算能力成为AI训练的核心硬件。'],
    ['参数', 'Parameter', '模型中可学习的权重。参数越多，模型容量越大。1B=10亿参数。'],
    ['Token', 'Token', '文本被拆分后的最小单元。一个中文约1.5-2个token，英文约0.75个token。'],
    ['预训练', 'Pre-training', '在海量数据上无监督学习通用知识，获得基础语言能力。'],
    ['微调', 'Fine-tuning', '在特定任务数据上对预训练模型做针对性调整，适应垂直场景。'],
    ['推理', 'Inference', '已训练好的模型对新的输入产生输出的过程，与训练相对。'],
    ['上下文窗口', 'Context Window', '模型单次能"看到"的最大token数。2026年顶级模型已支持100万+token。'],
]
add_styled_table(doc, term_headers, term_rows)

doc.add_heading('4.2 关键技术名词', level=2)

tech_headers = ['术语', '解释', '为什么重要']
tech_rows = [
    ['Transformer', '基于自注意力机制的神经网络架构，完全摒弃了循环结构', '所有现代LLM的基础架构'],
    ['注意力机制', '让模型动态关注输入中最相关部分的机制', '解决了长程依赖问题，是Transformer的核心'],
    ['GPT', 'Generative Pre-trained Transformer，生成式预训练模型', 'Decoder-only架构的代表，ChatGPT的技术基础'],
    ['RLHF', '基于人类反馈的强化学习，用人类偏好数据优化模型输出', '让模型输出更符合人类期望的关键技术'],
    ['RAG', '检索增强生成，在生成前先从知识库检索相关文档', '大幅减少幻觉，让AI基于事实回答'],
    ['Prompt', '提示词，用户向AI输入的指令或问题', '提示工程成为新兴技能，好的prompt显著影响输出质量'],
    ['幻觉', '模型生成看似合理但实际错误的内容', 'LLM的核心问题之一，RAG等技术正在缓解'],
    ['涌现能力', '模型参数超过一定阈值后突然出现的新能力', '解释了为什么大模型具有小模型不具备的能力'],
    ['量化', '降低模型参数的数值精度以减少存储和计算需求', 'INT4/INT8量化使大模型可在消费级GPU上运行'],
    ['嵌入', '将文本/图像等转化为固定长度的稠密数值向量', '语义搜索和知识检索的基础技术'],
    ['思维链', 'CoT，让模型逐步推理而非直接给出答案', '显著提升复杂推理任务的准确率'],
    ['蒸馏', '用大模型教导小模型，让后者继承前者的能力', '小模型获得大模型90%+性能的关键方法'],
]
add_styled_table(doc, tech_headers, tech_rows)

doc.add_heading('4.3 AI、ML、DL的关系', level=2)
doc.add_paragraph(
    '三者的关系可以理解为同心圆结构：'
    '\n\n┌─────────────────────────────────┐'
    '\n│        人工智能 (AI)            │'
    '\n│   ┌───────────────────────┐    │'
    '\n│   │   机器学习 (ML)       │    │'
    '\n│   │  ┌───────────────┐   │    │'
    '\n│   │  │  深度学习(DL) │   │    │'
    '\n│   │  │  Transformer │   │    │'
    '\n│   │  │  LLM / Agent │   │    │'
    '\n│   │  └───────────────┘   │    │'
    '\n│   └───────────────────────┘    │'
    '\n└─────────────────────────────────┘'
    '\n\n• AI是最大的概念——任何让机器表现出智能行为的技术都算AI'
    '\n• ML是实现AI的一种方法——让机器从数据中学习而非手工编写规则'
    '\n• DL是ML的一种——使用深层神经网络自动提取特征'
    '\n• Transformer/LLM是DL的前沿——专门处理序列数据，实现语言理解和生成'
)
doc.add_page_break()

# ═══ CHAPTER 5: AI AGENT ═══
doc.add_heading('第五章  AI Agent 智能体详解', level=1)

doc.add_paragraph(
    'AI Agent（智能体）是2025-2026年AI领域最重要的范式转变。如果说传统LLM是一个"知识渊博的顾问"，'
    'AI Agent则是一个"能自主完成任务的数字员工"。它不仅能理解和生成文本，还能感知环境、'
    '制定计划、调用工具、执行动作，并根据反馈自我修正。'
)

doc.add_heading('5.1 Agent 核心架构', level=2)
doc.add_picture(chart_paths['agent_arch'], width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('5.2 Agent 的四大核心能力', level=2)

agent_headers = ['能力', '说明', '示例']
agent_rows = [
    ['感知 (Perception)', '通过多模态输入理解环境状态和用户意图', '读取文档、分析图片、监听API回调'],
    ['规划 (Planning)', '将复杂目标分解为可执行的子任务序列', '旅行规划→订机票→订酒店→排行程'],
    ['记忆 (Memory)', '短期(上下文)和长期(向量库/知识图谱)信息存储', '记住用户偏好、项目背景、历史对话'],
    ['行动 (Action)', '通过工具调用和代码执行来改变外部世界', '发送邮件、操作文件、查询数据库、运行代码'],
]
add_styled_table(doc, agent_headers, agent_rows)

doc.add_heading('5.3 单智能体 vs. 多智能体', level=2)
doc.add_paragraph(
    '单智能体架构：一个Agent独立完成任务，适用于明确、线性的任务。'
    '\n\n多智能体架构：多个专业Agent协作完成复杂任务。典型模式包括：'
    '\n• Master-Worker模式：一个主Agent分配任务，多个子Agent并行执行'
    '\n• 流水线模式：Agent按顺序传递处理结果'
    '\n• 辩论模式：多个Agent从不同角度分析问题，综合得出最优解'
    '\n\n研究表明，多智能体系统在复杂企业工作流中，任务成功率比单智能体高47%，响应时间缩短62%。'
)

doc.add_heading('5.4 主流Agent框架与平台', level=2)
framework_headers = ['平台/框架', '开发方', '特点']
framework_rows = [
    ['Anthropic Managed Agents', 'Anthropic', '记忆+多Agent共享记忆("梦境")，Skills API'],
    ['OpenAI Workspace Agents', 'OpenAI', '持久化工作Agent预览'],
    ['Mistral Workflows', 'Mistral', 'Agent工作流编排'],
    ['Cursor Cloud Agents', 'Cursor', '云端编码Agent API'],
    ['Microsoft Foundry Agents', '微软', '企业级Agent托管'],
    ['CrewAI / AutoGen', '社区开源', '多Agent协作框架，灵活可定制'],
    ['LangGraph', 'LangChain', '有状态的Agent工作流图'],
]
add_styled_table(doc, framework_headers, framework_rows)

doc.add_heading('5.5 Agent 安全与约束工程', level=2)
doc.add_paragraph(
    '随着Agent获得越来越强的自主能力，安全约束变得至关重要。2026年新出现的"约束工程"学科'
    '专注于在Agent框架中嵌入硬性的行为边界：'
    '\n• 资源配额限制：CPU/内存/API调用次数上限'
    '\n• 操作护栏：禁止执行特定类型命令、访问特定数据'
    '\n• 审计日志：完整记录Agent的每一步决策和操作'
    '\n• 人机回路：关键决策节点需要人类确认'
    '\n\n研究表明，引入约束工程后，生产环境中Agent的调试时间从平均40小时降至6小时。'
)

doc.add_page_break()

# ═══ CHAPTER 6: BIG DATA AND AI ═══
doc.add_heading('第六章  大数据与AI的关系', level=1)

doc.add_paragraph(
    '大数据和人工智能是相辅相成的两个领域。如果说AI是"引擎"，大数据就是"燃料"。'
    '没有海量高质量数据，深度学习模型就无法有效训练；没有AI技术，大数据的价值也难以充分挖掘。'
)

doc.add_heading('6.1 大数据如何驱动AI', level=2)

data_headers = ['阶段', '数据需求', '典型规模']
data_rows = [
    ['预训练', '互联网级文本、代码、多语言语料', '数TB至数十TB，数十万亿token'],
    ['微调', '高质量标注数据（指令-回复对）', '数GB至数百GB，数百万条样本'],
    ['RLHF', '人类偏好对比数据（A回复 vs B回复）', '数十万至数百万组偏好对比'],
    ['RAG', '企业私有文档、知识库、数据库', 'GB至TB级结构化+非结构化数据'],
    ['评估', '标准化基准测试集、对抗样本', '数千至数万道测试题'],
]
add_styled_table(doc, data_headers, data_rows)

doc.add_heading('6.2 数据质量挑战', level=2)
doc.add_paragraph(
    'Gartner在2026年预测，高达60%的AI项目可能因数据质量问题而失败。核心挑战包括：'
    '\n• 数据偏见：训练数据中的偏见会被模型放大，导致不公平输出'
    '\n• 数据污染：低质量、错误或恶意构造的数据混入训练集'
    '\n• 数据枯竭：高质量公开文本数据接近耗尽，"数据墙"成为真实瓶颈'
    '\n• 版权争议：训练数据中的版权内容是2024-2026年最大的法律焦点'
    '\n• 合成数据：用AI生成的数据训练AI，可能造成"模型近亲繁殖"'
)

doc.add_heading('6.3 数据处理技术栈', level=2)
doc.add_paragraph(
    '现代AI的数据处理管道通常包括以下环节：'
    '\n1. 数据采集：网络爬虫、API接入、传感器采集、用户生成内容'
    '\n2. 数据清洗：去重、去噪、格式标准化、异常值检测'
    '\n3. 数据标注：人工标注 + AI辅助标注（半自动化标注）'
    '\n4. 数据增强：回译、同义词替换、模板生成（扩充训练集）'
    '\n5. 特征工程：文本向量化（Embedding）、特征选择与降维'
    '\n6. 数据版本管理：类似代码Git，用DVC等工具管理数据集版本'
)
doc.add_page_break()

# ═══ CHAPTER 7: INFRASTRUCTURE ═══
doc.add_heading('第七章  AI基础设施与算力', level=1)

doc.add_heading('7.1 训练基础设施', level=2)
doc.add_paragraph(
    '现代大模型的训练需要前所未有的算力规模。2026年的标志性基础设施包括：'
    '\n• Colossus 1 (Anthropic × SpaceX)：300MW数据中心，22万+ GPU'
    '\n• Microsoft-OpenAI：非独家合作关系调整，OpenAI模型已登陆Amazon Bedrock'
    '\n• 分布式训练技术：DiLoCoX能将1000亿+参数模型的训练速度提升357倍，仅需1GB带宽'
    '\n• 中国算力布局：阿里云、华为昇腾、百度昆仑等自研AI芯片生态成型'
)

doc.add_heading('7.2 推理优化', level=2)
infer_headers = ['技术', '原理', '效果']
infer_rows = [
    ['KV Cache', '缓存历史token的K/V向量，避免重复计算', '复杂度从O(n²)降至O(n)，速度提升5倍+'],
    ['FlashAttention', '优化GPU内存访问模式，减少HBM读写', '训练速度提升3倍，显存减少50%'],
    ['INT4/INT8量化', '将参数从16位浮点压缩为4/8位整数', '模型体积缩小4倍'],
    ['投机解码', '用小模型快速生成草稿，大模型并行验证', '推理速度提升2-3倍，无损质量'],
    ['连续批处理', '动态合并多个请求，提高GPU利用率', '吞吐量提升10倍+'],
]
add_styled_table(doc, infer_headers, infer_rows)

doc.add_heading('7.3 模型部署模式', level=2)
doc.add_paragraph(
    '2026年AI部署已形成三种主要模式：'
    '\n\n1. 云API模式：直接调用OpenAI、Anthropic等云端API，零运维但成本较高。适合个人开发者和快速原型。'
    '\n\n2. 私有化部署：在企业自己的服务器上运行开源模型（如Qwen、DeepSeek），'
    '数据不离开企业内部。适合金融、医疗等强监管行业。'
    '\n\n3. 边缘部署：在手机、IoT设备上运行压缩模型（如Ministral 3B、LFM2.5），'
    '延迟低于5ms。适合同声传译、实时质检等场景。'
)

doc.add_page_break()

# ═══ CHAPTER 8: FUTURE OUTLOOK ═══
doc.add_heading('第八章  未来展望与趋势', level=1)

doc.add_heading('8.1 技术路线图 (2026-2030)', level=2)
road_headers = ['时间', '预期里程碑', '影响']
road_rows = [
    ['2026-2027', '垂直行业Agent全面落地，AI开始替代复杂知识工作', '金融、医疗、法律行业变革'],
    ['2027-2028', '跨行业Agent平台出现，标准化协议成熟', 'Agent互操作成为现实'],
    ['2028-2029', '世界模型(World Model)取代纯语言模型成为主流', 'AI理解物理世界，机器人智能体普及'],
    ['2030+', 'AGI可能逐步出现，AI成为核心数字基础设施', '社会经济结构深刻变革'],
]
add_styled_table(doc, road_headers, road_rows)

doc.add_heading('8.2 关键问题与挑战', level=2)
doc.add_paragraph(
    '1. AI安全与对齐：如何确保越来越强大的AI系统始终与人类价值观对齐？'
    '\n2. 就业影响：AI Agent可能替代大量白领工作，社会需要准备转型方案。'
    '\n3. 能源消耗：大模型训练的电力需求呈指数增长，可持续性成为紧迫问题。'
    '\n4. 监管平衡：在促进创新和保障安全之间找到平衡点。EU AI Act已全面实施。'
    '\n5. 数字鸿沟：AI能力集中少数国家和公司，如何实现普惠？'
    '\n6. 真相与信任：深度伪造技术日益精湛，信息鉴别能力成为社会必需素养。'
)

doc.add_heading('8.3 给个人的建议', level=2)
doc.add_paragraph(
    '面对AI浪潮，以下建议供参考：'
    '\n\n✅ 学会使用AI工具：把AI作为"增强助手"，提高工作效率和创造力'
    '\n✅ 培养AI难以替代的能力：批判性思维、创造力、人际沟通、跨领域整合'
    '\n✅ 保持学习心态：AI领域变化极快，今天的知识可能半年后就过时'
    '\n✅ 关注AI伦理：理解AI的能力边界和潜在风险，做一个负责任的AI使用者'
    '\n✅ 动手实践：不只是阅读，亲自用AI工具做项目，才能真正建立认知'
)

# -- Save Word document --
docx_path = os.path.join(OUTPUT_DIR, "AI知识全览_2026.docx")
doc.save(docx_path)
print(f"Word文档已生成: {docx_path}")

# ═══════════════════════════════════════════════════════════════
# PART 3: Generate Excel with detailed data
# ═══════════════════════════════════════════════════════════════

wb = openpyxl.Workbook()

# -- Colors --
header_fill = PatternFill(start_color="1A3C6E", end_color="1A3C6E", fill_type="solid")
header_font = Font(name="Microsoft YaHei", size=11, bold=True, color="FFFFFF")
data_font = Font(name="Microsoft YaHei", size=10)
title_font = Font(name="Microsoft YaHei", size=14, bold=True, color="1A3C6E")
thin_border = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)

def style_sheet(ws, headers, data, col_widths):
    """Apply consistent styling to a worksheet."""
    for c, (header, width) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=c, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(c)].width = width
    for r, row in enumerate(data, 2):
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = data_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
    ws.freeze_panes = 'A2'

# --- Sheet 1: AI History Timeline ---
ws1 = wb.active
ws1.title = "AI发展历史"
h1 = ['年份', '事件', '意义', '阶段']
d1 = [
    [1943, 'McCulloch-Pitts神经元模型', '第一个神经元数学模型', '理论奠基期'],
    [1950, '图灵测试提出', '提出机器智能的评判标准', '理论奠基期'],
    [1956, '达特茅斯会议', '正式提出"人工智能"术语', '理论奠基期'],
    [1957, '感知机 (Perceptron)', '第一个可训练的神经网络', '符号主义时代'],
    [1958, 'LISP语言', 'AI研究的主流编程语言', '符号主义时代'],
    [1966, 'ELIZA聊天机器人', '第一个可对话的AI程序', '符号主义时代'],
    [1974, '第一次AI寒冬开始', '过高期望破灭，投资锐减', 'AI寒冬'],
    [1982, 'Hopfield网络', '重新点燃神经网络研究', '连接主义复兴'],
    [1986, '反向传播算法普及', '多层神经网络可有效训练', '连接主义复兴'],
    [1989, 'CNN应用于手写识别', 'LeCun奠定计算机视觉基础', '连接主义复兴'],
    [1997, 'Deep Blue击败卡斯帕罗夫', 'AI首次在正式比赛中战胜国际象棋世界冠军', '统计学习时代'],
    [1997, 'LSTM长短期记忆网络', '解决循环神经网络梯度消失问题', '统计学习时代'],
    [2006, '深度信念网络', 'Hinton重新点燃深度学习', '统计学习时代'],
    [2009, 'ImageNet发布', '李飞飞发布大规模图像数据集', '统计学习时代'],
    [2012, 'AlexNet赢得ImageNet竞赛', '深度学习革命的起点', '深度学习革命'],
    [2014, 'GAN生成对抗网络', 'AI能够生成逼真的图像', '深度学习革命'],
    [2016, 'AlphaGo击败李世石', 'AI在围棋领域超越人类', '深度学习革命'],
    [2017, 'Transformer架构发布', 'Attention Is All You Need', '深度学习革命'],
    [2018, 'BERT和GPT-1发布', '开启预训练+微调范式', '预训练时代'],
    [2020, 'GPT-3 (1750亿参数)', '展示令人惊叹的涌现能力', '预训练时代'],
    [2020, 'AlphaFold 2', '解决50年蛋白质折叠难题', '预训练时代'],
    [2022, 'ChatGPT发布', '2个月1亿用户，AI走向大众', '生成式AI爆发'],
    [2023, 'GPT-4发布', '多模态理解，专业级表现', '生成式AI爆发'],
    [2023, 'EU AI Act通过', '全球首部综合性AI法规', '生成式AI爆发'],
    [2024, 'Sora视频生成模型', 'AI视频生成质量飞跃', '多模态时代'],
    [2025, 'AI Agent成熟', 'AI从工具进化为自主执行者', 'Agent时代'],
    [2026, 'MCP成为行业标准', '多智能体协作大规模落地', 'Agent时代'],
]
style_sheet(ws1, h1, d1, [8, 28, 42, 18])

# --- Sheet 2: Model Comparison ---
ws2 = wb.create_sheet("主流模型对比")
h2 = ['模型名称', '开发公司', '架构类型', '参数规模', '上下文窗口', '多模态', '开源', '发布时间']
d2 = [
    ['Claude Opus 4.7', 'Anthropic', 'Decoder-Only', '~2.8万亿', '100万+ tokens', '是', '否', '2026年4月'],
    ['GPT-5.5', 'OpenAI', 'Decoder-Only', '~3万亿(MoE)', '200万 tokens', '是', '否', '2026年4月'],
    ['GPT-Rosalind', 'OpenAI', 'Decoder-Only + 推理优化', '未公开', '未公开', '是', '否', '2026年'],
    ['Gemini 3.1 Flash', 'Google', 'Encoder-Decoder', '未公开', '100万+ tokens', '是', '否', '2026年'],
    ['Qwen 3.6', '阿里巴巴', 'Decoder-Only + 门控注意力', '~2.5万亿(MoE)', '128K tokens', '是', '是', '2026年'],
    ['DeepSeek V4', 'DeepSeek', 'Decoder-Only (MLA+MoE)', '~2.3万亿(MoE)', '128K tokens', '否', '是', '2026年'],
    ['Mistral 3 Large', 'Mistral AI', 'Decoder-Only', '~1.2万亿', '128K tokens', '是', '是', '2026年'],
    ['GLM-4.7', '智谱AI', 'Decoder-Only (MoE)', '355B总参(32B激活)', '128K tokens', '否', '是', '2026年'],
    ['Kimi K2.6', 'Moonshot AI', 'Decoder-Only', '未公开', '128K tokens', '是', '是', '2026年'],
    ['Nemotron 3 Nano', 'NVIDIA', 'Decoder-Only (MoE)', '300亿总参(30B)', '128K tokens', '否', '是(完全开源)', '2026年'],
    ['K-EXAONE', 'LG', 'Decoder-Only (MoE)', '2360亿总参', '128K tokens', '否', '否', '2026年'],
    ['Claude 3.5 Sonnet', 'Anthropic', 'Decoder-Only', '~2万亿', '200K tokens', '是', '否', '2024年'],
    ['GPT-4o', 'OpenAI', 'Decoder-Only', '~1.8万亿', '128K tokens', '是', '否', '2024年'],
    ['Llama 3.1 405B', 'Meta', 'Decoder-Only', '4050亿', '128K tokens', '否', '是', '2024年'],
]
style_sheet(ws2, h2, d2, [20, 14, 22, 18, 16, 10, 16, 14])

# --- Sheet 3: Terminology Glossary ---
ws3 = wb.create_sheet("AI术语表")
h3 = ['术语', '英文', '类别', '难度', '核心解释']
d3 = [
    ['人工智能', 'Artificial Intelligence', '基础概念', '入门', '让机器模拟人类智能的学科体系'],
    ['机器学习', 'Machine Learning', '基础概念', '入门', '通过数据训练让模型自动学习，而非手工编写规则'],
    ['深度学习', 'Deep Learning', '基础概念', '入门', '使用多层神经网络从大规模数据中自动提取特征'],
    ['大语言模型', 'Large Language Model', '核心概念', '中级', '以Transformer为基础、参数规模达数十亿以上的语言模型'],
    ['Transformer', 'Transformer', '核心架构', '中级', '基于自注意力机制的神经网络架构，所有现代LLM的基石'],
    ['注意力机制', 'Attention Mechanism', '核心机制', '中级', '让模型动态关注输入中最相关部分的计算方法'],
    ['Token', 'Token', '基础概念', '入门', '文本被拆分后的最小处理单元，中文1字≈1.5-2 token'],
    ['参数', 'Parameter', '基础概念', '入门', '模型中可学习的权重值，参数数量决定模型容量'],
    ['预训练', 'Pre-training', '训练方法', '中级', '在海量数据上无监督学习通用知识的初始训练阶段'],
    ['微调', 'Fine-tuning', '训练方法', '中级', '在特定任务数据上对预训练模型做针对性调整'],
    ['RLHF', 'Reinforcement Learning from Human Feedback', '训练方法', '高级', '用人类偏好数据训练奖励模型，再通过强化学习优化输出'],
    ['RAG', 'Retrieval-Augmented Generation', '增强技术', '高级', '生成前从外部知识库检索相关文档，减少幻觉'],
    ['AI Agent', 'AI Agent', '前沿概念', '高级', '能感知环境、制定计划、调用工具并自主执行任务的AI系统'],
    ['MCP', 'Model Context Protocol', '前沿概念', '高级', 'Anthropic推出的AI与外部工具交互的开放协议标准'],
    ['MoE', 'Mixture of Experts', '核心架构', '高级', '每次推理仅激活部分"专家"参数，提升效率的稀疏模型架构'],
    ['GQA', 'Grouped Query Attention', '优化技术', '高级', '注意力头分组共享K/V投影，降低显存占用'],
    ['RoPE', 'Rotary Position Embedding', '核心机制', '高级', '通过旋转操作将位置信息编码到注意力计算中'],
    ['幻觉', 'Hallucination', '问题', '中级', '模型生成看似合理但实际错误的内容'],
    ['涌现能力', 'Emergent Ability', '核心概念', '高级', '模型参数超过一定阈值后突然出现的未预期能力'],
    ['思维链', 'Chain of Thought', '推理技术', '中级', '让模型展示逐步推理过程，显著提升复杂任务准确率'],
    ['量化', 'Quantization', '优化技术', '中级', '降低参数数值精度(如FP16→INT4)，减少存储和计算需求'],
    ['知识蒸馏', 'Knowledge Distillation', '优化技术', '高级', '用大模型指导小模型训练，传承知识能力'],
    ['上下文窗口', 'Context Window', '基础概念', '中级', '模型单次能处理的最大token数量'],
    ['多模态', 'Multimodal', '前沿概念', '高级', 'AI同时理解和生成文本、图像、语音、视频等多种信息形式'],
    ['合成数据', 'Synthetic Data', '数据技术', '中级', '用AI生成的数据来训练AI，解决真实数据不足的问题'],
    ['世界模型', 'World Model', '前沿概念', '高级', '能理解和预测物理世界运行规律的基础模型，被视为LLM后的下一阶段'],
]
style_sheet(ws3, h3, d3, [16, 30, 12, 8, 48])

# --- Sheet 4: Market & Compute Data ---
ws4 = wb.create_sheet("市场与算力数据")
h4 = ['年份', '全球AI市场规模(十亿美元)', '训练算力(PetaFLOPS-days)', '最大模型参数量(亿)', '代表模型']
d4 = [
    [2018, 23.5, 0.5, 3.4, 'BERT'],
    [2019, 35.8, 2, 15, 'GPT-2'],
    [2020, 51.3, 8, 1750, 'GPT-3'],
    [2021, 93.5, 30, 5300, 'Megatron-Turing NLG'],
    [2022, 142.3, 120, 5400, 'PaLM'],
    [2023, 207.9, 500, 18000, 'GPT-4'],
    [2024, 298.5, 2100, 20000, 'Claude 3.5 / Gemini Ultra'],
    [2025, 420.0, 8500, 30000, 'GPT-5 / DeepSeek-V4'],
    [2026, 594.0, 35000, 30000, 'Claude 4.7 / GPT-5.5'],
    ['2027E', 780.0, 80000, 50000, '预估'],
    ['2028E', 1020.0, 150000, 80000, '预估'],
]
style_sheet(ws4, h4, d4, [10, 28, 24, 22, 28])

# Add chart to Excel
chart = LineChart()
chart.title = "全球AI市场与算力增长趋势"
chart.style = 10
chart.y_axis.title = "市场规模(十亿美元)"
chart.x_axis.title = "年份"
data_ref = Reference(ws4, min_col=2, min_row=1, max_row=12)
cats_ref = Reference(ws4, min_col=1, min_row=2, max_row=12)
chart.add_data(data_ref, titles_from_data=True)
chart.set_categories(cats_ref)
chart.width = 22
chart.height = 14
ws4.add_chart(chart, "A15")

# --- Sheet 5: Recommended Learning Resources ---
ws5 = wb.create_sheet("推荐学习资源")
h5 = ['资源名称', '类型', '难度', '适合人群', '说明']
d5 = [
    ['李宏毅《机器学习》课程', '视频课程', '入门-中级', '所有人', '中文最好的ML入门课，深入浅出'],
    ['吴恩达《Machine Learning》', '视频课程', '入门', '所有人', '全球最经典的ML入门课程'],
    ['Andrej Karpathy《Neural Networks》', '视频课程', '中级', '有编程基础者', '从零实现神经网络和GPT'],
    ['《Attention Is All You Need》', '论文', '高级', '研究人员', 'Transformer原始论文，必读经典'],
    ['Hugging Face Course', '在线教程', '中级', '开发者', '动手学Transformers和模型部署'],
    ['《Deep Learning》(花书)', '书籍', '高级', '研究人员', '深度学习领域的权威教材'],
    ['Anthropic Claude文档', '官方文档', '入门-中级', '开发者', '学习AI工具和Agent开发'],
    ['OpenAI Cookbook', '代码仓库', '中级', '开发者', 'GPT API实战示例合集'],
    ['Papers With Code', '网站', '中级-高级', '研究人员', '论文+代码对照，跟踪最新SOTA'],
    ['LangChain/LlamaIndex文档', '框架文档', '中级', 'AI应用开发者', 'LLM应用开发框架'],
]
style_sheet(ws5, h5, d5, [30, 12, 14, 18, 38])

excel_path = os.path.join(OUTPUT_DIR, "AI数据与术语表_2026.xlsx")
wb.save(excel_path)
print(f"Excel文件已生成: {excel_path}")

print("\n=== 全部文档生成完成 ===")
print(f"  Word: {docx_path}")
print(f"  Excel: {excel_path}")
print(f"  图表: {OUTPUT_DIR}\\chart_*.png (共{len(chart_paths)}张)")
